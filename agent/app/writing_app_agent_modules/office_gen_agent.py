import ast
import json
import traceback

from agent.agent_node import AgentNode
from agent.model.DeepseekModel import llm


class OfficeGenAgent(AgentNode):

    def __init__(self):
        super().__init__()
        self.llm = llm
        self.toolData = ""
        self.max_step = 8
        self.current_step = 0

    def get_prompt(self, chat_no=None):
        content = f"""
        你是一名公文提取助手，你的目标提取用户问题中信息。

        基本信息:
            {self.get_context_data(chat_no, "baseInfo")}

        你必须严格按照如下步骤流程执行:
            1.理解用户意图
            2.根据用户提供的内容，提取相关信息（参考标题/参考内容/公文类型）
        
        
        公文类型:
            条例、规定、办法、决定、命令、指示、批复、通知、通报、公告、通告、议案、请示、报告、 调查报告 、 总结报告 、函、 会议纪要
        
        返回json格式的数据:
          {{"code":0,"reply":<string>,"reference_content":<string>,"official_type":<string>,"reference_title":<string>}}

        返回值说明:
          - `code`： 0 -> 执行成功, 1 -> 执行失败
          - `reply`: 字符串，成功/失败
          - `reference_content`: 参考内容
          - `reference_title`: 参考标题
          - `official_type`: 公文类型
         """
        message = {"role": "system", "content": content}
        return message

    def query_desc(self) -> str:
        desc = """
         ArticleGenAgent -> 这是一个用于公文生成的Agent
         """
        return desc

    def query_name(self) -> str:
        return "ArticleGenAgent"

    async def action(self, chat_no: str, json_data: dict):
        if json_data["code"] != 0:
            return json_data
        result_outline = await self.to_outline(chat_no,json_data["reference_content"],json_data["official_type"],json_data["reference_title"])
        if result_outline["code"] != 0:
            return result_outline
        result_content = await self.to_content(chat_no,json_data["reference_title"],json_data["official_type"],json_data["reference_content"],result_outline["reply"])
        if result_content["code"] == 0:
            await self.format_word(chat_no,result_content["reply"])
        return result_content

    async def to_summary(self,chat_no: str,reference_content: str,official_type: str,reference_title: str):
        content = f"""
                    你是一名公文摘要写作助手，你的目标根据上下信息生成公文摘要。

                    基本信息:
                    {self.get_context_data(chat_no, "baseInfo")}

                    你必须严格按照如下步骤流程执行:
                        1.理解用户意图
                        2.根据上下文信息生成公文摘要

                    上下文信息:
                        1.参考标题
                        {reference_title}

                        2.公文类型
                        {official_type}

                        3.参考内容
                        {reference_content}

                    返回json格式的数据:
                    {{"code":0,"reply":<string>}}

                    返回值说明:
                    - `code`： 0 -> 执行成功, 1 -> 执行失败
                    - `reply`: 字符串，公文摘要
                """
        system_message = {"role": "system", "content": content}
        result = await llm.acall(
            json.dumps([system_message, {"role": "user", "content": "根据上下文信息生成公文摘要"}]))
        response_data = result["choices"][0]["message"]["content"]
        return json.loads(response_data)
    async def to_outline(self,chat_no: str,reference_content: str,official_type: str,reference_title: str):
        content = f"""
            你是一名公文目录写作助手，你的目标根据上下信息生成公文目录。
            
            基本信息:
            {self.get_context_data(chat_no, "baseInfo")}
            
            你必须严格按照如下步骤流程执行:
                1.理解用户意图
                2.根据上下文信息一步一步的构建目录模版（各级标题及子标题）
                
            上下文信息:
                1.参考标题
                {reference_title}
                
                2.公文类型
                {official_type}
                
                3.参考内容
                {reference_content}
            
            返回json格式的数据:
            {{"code":0,"reply":<string>}}

            返回值说明:
            - `code`： 0 -> 执行成功, 1 -> 执行失败
            - `reply`: 字符串，内容目录
        """
        system_message = {"role": "system", "content": content}
        result = await llm.acall(json.dumps([system_message,{"role": "user","content": "根据上下文信息生成公文目录"}]))
        response_data = result["choices"][0]["message"]["content"]
        return json.loads(response_data)

    async def to_content(self,chat_no: str,reference_title: str,official_type: str,reference_content: str,official_outline: str):
        content = f"""
            你是专业的党政机关公文写作助手，任务是基于提供的上下文信息生成结构完整、内容详实、细节丰富、格式规范的正式公文。请严格按照以下要求操作：

            基本信息：
                任务目标：根据给定的参考目录，结合参考内容，填充详实的内容细节，生成规范的公文正文。
                字数要求：输出内容不少于1800字。
                格式要求：必须符合国家党政机关公文格式规范，包括版头、主体、版记三部分。
                内容要求：结构完整，逻辑清晰，内容细节充分，不得遗漏参考内容中的任何要素，如无细节需合理补充。
                风格要求：语言正式、严谨、准确，兼顾可读性和权威性。
                输出格式：标准 JSON 格式输出。

           你必须严格按照如下步骤执行：
                1.理解用户意图与任务背景：明确参考标题、公文类型及核心任务目标。
                2.结构填充并扩展内容：
                3.根据“参考目录”填充正文结构；
                    在每一节中结合“参考内容”展开细节；
                    若参考内容缺乏必要细节，请自行合理补充背景、过程、数据、案例、结果等信息。
                    正文内容应细致展开，例如：
                        涉及事件时要交代发生背景、参与方、关键节点、具体做法、时间安排、遇到的问题与解决方案、成果评估等；
                        涉及到专有名词或术语或关键信息或文章中特指的名称时，应该给出详细说明
                        若有多个措施或选项，需分别列出并对比利弊；
                        说明原因时应结合数据、案例、政策依据。
                4.格式化处理：
                    按照党政机关公文格式排版；
                    各要素位置、字体、标点、符号、单位均应规范；
                    主体内容标题居中，段落首行缩进2字符；
                    层级结构使用“一、”“（一）”“1.”“（1）”等格式标注。
                5.生成最终内容：


           上下文信息:

                1. 补充信息
                {self.get_context_data(chat_no, "baseInfo")}

                2.参考标题，需作为公文标题核心内容展开
                {reference_title}

                3.公文类型，如通知、通报、请示、报告等，决定文风与格式细节
                {official_type}

                4.参考内容，提供的初步内容参考，不得遗漏其中的任何关键信息。
                {reference_content}

                5.参考目录，是正文内容结构依据，应完整匹配且内容充分展开。
                {official_outline}

           返回json格式的数据:
            {{"code":0,"reply":<string>}}

           返回值说明:
            - `code`： 0 -> 执行成功, 1 -> 执行失败
            - `reply`: 字符串，公文内容
        """

        system_message = {"role": "system", "content": content}
        result = await llm.acall(
            json.dumps([system_message, {"role": "user", "content": "生成公文内容"}]))
        response_data = result["choices"][0]["message"]["content"]
        return json.loads(response_data)
        
    async def format_word(self,chat_no: str,official_content: str):
        content = f"""
                你的任务是：将输入的党政机关公文内容，严格按照《党政机关公文格式》（GB/T 9704—2012）要求写入 Word 文档。

                你必须严格依照以下步骤进行操作：
                    1.你不能改变公文内容，只能做格式调整
                    2.调用工具将输入的公文内容写入 Word 文档（必须符合公文排版要求）；
                    3.必须确保所有公文内容完整准确地写入文档；
                    
                可调用的工具:
                  - `word_exec(code: str)`: 通过执行python脚本代码将公文写入word文档，返回word文档地址
                        - `code`: 写入word文档的python脚本代码
                  
                Python 脚本规范（传入 word_exec 的参数）：
                    1.执行结果必须赋值给变量 result；
                    2.处理中文防止乱码；
                    3.设置正确的字体、字号、行距、段落缩进和对齐方式；
                    4.返回文件保存路径格式如下：
                        本地路径前缀：static/temp/image/{chat_no}/（路径不存在时需自动创建）
                        访问地址前缀：http://127.0.0.1:8000/static/temp/image/{chat_no}/
                    5.字体设置示例如下：
                        
                    from docx import Document
                    from docx.shared import Pt
                    from docx.oxml.ns import qn
                    from docx.shared import RGBColor
                    import os
                    
                    # 创建文档对象
                    doc = Document()

                    # 设置页面为A4，页边距符合GB/T 9704-2012标准（单位：厘米）
                    section = doc.sections[0]
                    section.page_height = Cm(29.7)
                    section.page_width = Cm(21.0)
                    section.top_margin = Cm(3.7)
                    section.bottom_margin = Cm(2.5)
                    section.left_margin = Cm(2.8)
                    section.right_margin = Cm(2.2)
                    
                    # 定义通用样式函数,你可以通过以下方法设置字体名称，大小，颜色等信息
                    def set_font(paragraph, font_name, font_size_pt, bold=False, color=None):
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run.font.size = Pt(font_size_pt)
                        run.bold = bold
                        run.font.color.rgb = RGBColor(*(color if color else (0, 0, 0)))  # 默认黑色
                        return paragraph
                    
                    # 使用set_font函数设置字体样式
                    
                    # 标题内容格式
                    p = doc.add_paragraph('标题内容' )
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font(p, font_name='仿宋', size=22)
                    
                    # 正文内容格式
                    p = doc.add_paragraph('正文内容')
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    set_font(p, '仿宋', 16)

                    # 成文日期内容格式
                    p = doc.add_paragraph('成文日期内容')
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    set_font(p, font_name='仿宋', size=16)

                公文排版要求摘要（GB/T 9704-2012）：
                    纸张尺寸: A4（210mm × 297mm）
                    页边距: 上 3.7cm、下 2.5cm、左 2.8cm、右 2.2cm
                    标题字体字号: 2号小标宋体，居中排布
                    正文字体字号: 3号仿宋体，首行缩进2字符，固定行距20磅
                    成文日期: 小四仿宋体，右对齐
                    抄送机关:抄送机关
                    印发机关和印发日期: 4 号仿宋体字
                    附件与页码、版记: 按规范独立处理，附件另起页
                    字体颜色无特殊说明默认黑色
                    
                  
                公文内容:
                  {official_content}
                  
                  
                返回json格式:
                  {{"code":<int>,"reply":<string>,"tool_use": <bool>,"tool_name": <string>,"args": <list>}}

                返回值说明:
                    - `code`： 0 -> 执行成功, 1 -> 执行失败
                    - `reply`:  如果成功 -> 成功原因, 如果失败 -> 失败原因
                    - `tool_use`:  true -> 需使用工具, false -> 不使用工具
                    - `tool_name`: 使用的工具名称
                    - `args`: 工具所需的参数列表
                
                """
        system_message = {"role": "system", "content": content}
        result = await llm.acall(
            json.dumps([system_message, {"role": "user", "content": "将公文的所有内容写入文档"}]))
        response_data = result["choices"][0]["message"]["content"]
        json_data = json.loads(response_data)
        await self.word_exec(json_data["args"][0])


    async def word_exec(self, code: str):

        return await self.execute_with_imports(code, {})

    async def execute_with_imports(self, code, params):
        # 用同一个字典做全局和局部，保持import生效且能被调用
        local_vars = {}
        try:
            exec(code, local_vars, local_vars)  # 这里第二、第三个参数都用 local_vars
            if params:
                local_vars.update(params)
            return {"code": 0, 'result': local_vars.get('result', None), 'error': None}
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {"code": -1, 'result': None, 'error': str(e)}


instance = OfficeGenAgent()

