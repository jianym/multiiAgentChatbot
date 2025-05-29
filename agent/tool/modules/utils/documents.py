import io

import requests
from bs4 import BeautifulSoup
import pdfplumber
from io import BytesIO
from langchain.text_splitter import RecursiveCharacterTextSplitter
import chardet
from docx import Document
from playwright.async_api import async_playwright


def word_document(content, link):
    docs = []
    word_text = ""
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    for page in content.paragraphs:
        word_text += p.text
    word_text = " ".join(word_text.split())
    splitted_text = splitter.split_text(word_text)
    title = "PDF Document"
    # 创建文档对象并添加到列表
    for text in splitted_text:
        docs.append({"pageContent": text, "title": title, "url": link})
    return docs


def pdfDoucment(content, link):
    docs = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    # 处理 PDF 文件
    with pdfplumber.open(BytesIO(content)) as pdf:
        pdf_text = ""
        for page in pdf.pages:
            pdf_text += page.extract_text()

            # 清理文本
        pdf_text = " ".join(pdf_text.split())
        splitted_text = splitter.split_text(pdf_text)
        title = "PDF Document"

        # 创建文档对象并添加到列表
        for text in splitted_text:
            docs.append({"pageContent": text, "title": title, "url": link})
    return docs


def htmlDocument(content, link):
    docs = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    # 处理 HTML 文件
    soup = BeautifulSoup(content, 'html.parser')

    # 去除所有 <a> 标签及其 href 属性
    # for a_tag in soup.find_all('a'):
    #     a_tag.decompose()  # 去除 <a> 标签及其内容

    # 获取纯文本，去除所有标签
    html_text = soup.get_text(separator=' ', strip=True)

    html_text = " ".join(html_text.split())  # 去除多余空格

    # 使用文本拆分器进行拆分
    splitted_text = splitter.split_text(html_text)

    # 获取 HTML 标题
    title_tag = soup.find('title')
    title = title_tag.string if title_tag else link

    # 创建文档对象并添加到列表
    for text in splitted_text:
        docs.append({"pageContent": text, "title": title, "url": link})

    return docs


async def getDocumentsFromLinks(links):
    docs = []
    for link in links:
        # 确保链接以 http:// 或 https:// 开头
        if not link.startswith('http://') and not link.startswith('https://'):
            link = f'https://{link}'

        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br',
                'Referer': 'https://www.google.com'
            }
            cookies = {
                "x-waf-captcha-referer": "https://www.google.com",
            }
            res = requests.get(link, cookies=cookies, headers=headers)
            # res.encoding = 'utf-8'
            res.raise_for_status()  # 如果请求失败，抛出异常
            content_type = res.headers['Content-Type']
            # 使用 chardet 自动检测响应的编码
            detected_encoding = chardet.detect(res.content)['encoding']
            if 'pdf' in content_type:
                docs.extend(pdfDoucment(res.content, link))
            elif "wordprocessingml.document" in content_type or "application/msword" in content_type:
                doc_data = io.BytesIO(res.content)
                doc = Document(doc_data)
                docs.extend(word_document(doc, link))
            else:
                res.encoding = detected_encoding  # 设置正确的编码
                docs.extend(htmlDocument(res.text, link))

        except Exception as e:
            html = fetch_with_playwright(link)
            if html:
                docs.extend(htmlDocument(html, link))
            else:
                print(f"[fallback failed] {link}")
    return docs

# 抓取免费代理
async def fetch_free_proxies():
    url = "https://www.free-proxy-list.net/"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    }
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
    except Exception as e:
        print(f"抓取代理列表失败：{e}")
        return []

    soup = BeautifulSoup(resp.text, "html.parser")
    proxies = []
    table = soup.find("table", id="proxylisttable")
    if not table:
        print("找不到代理表格")
        return proxies

    rows = table.tbody.find_all("tr")
    for row in rows:
        cols = row.find_all("td")
        ip = cols[0].text.strip()
        port = cols[1].text.strip()
        https = cols[6].text.strip()
        scheme = "https" if https == "yes" else "http"
        proxy = f"{scheme}://{ip}:{port}"
        proxies.append(proxy)
    return proxies

# 测试代理有效性
async def test_proxy(proxy):
    proxies = {"http": proxy, "https": proxy}
    test_url = "https://httpbin.org/ip"
    try:
        r = requests.get(test_url, proxies=proxies, timeout=5)
        if r.status_code == 200:
            print(f"[有效] {proxy} 返回IP: {r.json().get('origin')}")
            return True
    except Exception as e:
        print(f"[无效] {proxy} 异常: {e}")
    return False

async def fetch_with_playwright(link):
    """使用 Playwright 浏览器获取页面 HTML 内容"""
    try:

        proxies = await fetch_free_proxies()
        print(f"抓取到 {len(proxies)} 个代理，开始测试有效性...")

        proxy_url = ""
        for proxy in proxies:
            if await test_proxy(proxy):
                proxy_url = proxy
                break
        print(f"可用代理IP{proxy_url}...")

        if len(proxy_url) <= 1:
            return ""
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            context = await browser.new_context(proxy={"server": proxy_url})
            page = await context.new_page()
            try:
                await page.goto(link, timeout=10000)
                return await page.content()
            except Exception as e:
                print(f"代理 {proxy_url} 访问失败: {e}")
            finally:
                await browser.close()
    except Exception as e:
        print(f"[Playwright] Failed to fetch {link}: {e}")
        return None


# def exec_js(html_content):
#     # 使用 BeautifulSoup 提取 JavaScript 代码
#     soup = BeautifulSoup(html_content, 'html.parser')
#
#     # 获取所有的 <script> 标签
#     scripts = soup.find_all('script')
#
#     # 逐个检查脚本内容，找到你需要的 JavaScript
#     for script in scripts:
#         script_content = script.string
#         if script_content:
#             # 创建 JavaScript 执行环境
#             ctx = execjs.compile(script_content)
#             print(ctx)


if __name__ == '__main__':
    docs = getDocumentsFromLinks(["https://www.maoyan.com/"])
    print(docs)
